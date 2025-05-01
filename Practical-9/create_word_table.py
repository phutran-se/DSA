from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_word_table(input_file, output_file):
    """ Create a Word document with a table from a text file
    Args:
        input_file (str): Path to input text file (e.g., sorting_runtimes.txt)
        output_file (str): Path to output Word file (e.g., Sorting_Runtimes.docx)
    """
    # Create a new Word document
    doc = Document()
    
    # Add a title
    doc.add_heading('Runtime Results of Sorting Algorithms', level=1)
    
    # Read data from input file
    with open(input_file, 'r') as f:
        lines = f.readlines()
    
    # Filter out separator lines and empty lines
    data = [line.strip().split('|') for line in lines if '|' in line and not line.startswith('-----')]
    data = [[cell.strip() for cell in row] for row in data]
    
    # Create table
    table = doc.add_table(rows=len(data), cols=4)
    table.style = 'Table Grid'  # Apply grid style
    
    # Populate table
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.cell(i, j).text = cell
            # Set font and alignment
            for paragraph in table.cell(i, j).paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                paragraph.alignment = 1  # Center alignment
    
    # Format header row
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Add shading (light gray)
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'D3D3D3')  # Light gray color
        tcPr.append(shd)
    
    # Save document
    doc.save(output_file)
    print(f"Word document saved as {output_file}")

if __name__ == "__main__":
    create_word_table('sorting_runtimes.md', 'Sorting_Runtimes.docx')